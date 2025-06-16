"""Word to Markdown conversion using multiple strategies."""

import logging
import subprocess
import tempfile
from pathlib import Path
from typing import Dict, Any, Optional
import pypandoc
import docx2python
from docx import Document


class WordToMarkdownConverter:
    """Convert Word documents to Markdown with multiple strategies."""

    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.markdown_options = config.get('markdown_options', {})
        self.logger = logging.getLogger(__name__)

    def convert(self, word_path: str, output_path: Optional[str] = None) -> str:
        """
        Convert Word document to Markdown.

        Args:
            word_path: Path to input Word document
            output_path: Path for output Markdown file (optional)

        Returns:
            Path to converted Markdown file
        """
        word_path = Path(word_path)

        if output_path is None:
            output_path = word_path.with_suffix('.md')
        else:
            output_path = Path(output_path)

        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Try different conversion strategies in order of preference
        strategies = [
            self._convert_with_pandoc,
            self._convert_with_docx2python,
            self._convert_with_python_docx
        ]

        for strategy in strategies:
            try:
                self.logger.info(
                    f"Trying {strategy.__name__} for {word_path.name}")
                result = strategy(str(word_path), str(output_path))
                if result and Path(result).exists():
                    self.logger.info(
                        f"Successfully converted using {strategy.__name__}")
                    return str(result)
            except Exception as e:
                self.logger.warning(f"{strategy.__name__} failed: {e}")
                continue

        raise RuntimeError(
            f"All Word to Markdown conversion strategies failed for {word_path}")

    def _convert_with_pandoc(self, word_path: str, output_path: str) -> str:
        """Convert using pypandoc (best for academic papers)."""
        try:
            # Configure pandoc options for academic papers
            extra_args = [
                '--extract-media=' +
                str(Path(output_path).parent /
                    self.markdown_options.get('media_dir', 'media')),
                '--wrap=none',  # Don't wrap long lines
                '--markdown-headings=atx',  # Use # style headings
            ]

            # Add table format option
            table_format = self.markdown_options.get('table_format', 'github')
            if table_format == 'grid':
                extra_args.append('--columns=120')

            # Convert with pypandoc
            output = pypandoc.convert_file(
                word_path,
                'markdown',
                outputfile=output_path,
                extra_args=extra_args
            )

            # Post-process the markdown for better formatting
            self._post_process_markdown(output_path)

            return output_path

        except Exception as e:
            raise RuntimeError(f"Pandoc conversion failed: {e}")

    def _convert_with_docx2python(self, word_path: str, output_path: str) -> str:
        """Convert using docx2python library."""
        try:
            # Extract content using docx2python
            content = docx2python.docx2python(word_path)

            # Build markdown from extracted content
            markdown_content = self._build_markdown_from_docx2python(content)

            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)

            return output_path

        except Exception as e:
            raise RuntimeError(f"docx2python conversion failed: {e}")

    def _convert_with_python_docx(self, word_path: str, output_path: str) -> str:
        """Convert using python-docx library (basic conversion)."""
        try:
            doc = Document(word_path)
            markdown_lines = []

            # Process document paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    markdown_lines.append("")
                    continue

                # Handle headings
                if paragraph.style.name.startswith('Heading'):
                    level = self._get_heading_level(paragraph.style.name)
                    markdown_lines.append(f"{'#' * level} {text}")
                    markdown_lines.append("")
                else:
                    # Regular paragraph
                    markdown_lines.append(text)
                    markdown_lines.append("")

            # Process tables
            for table in doc.tables:
                table_md = self._convert_table_to_markdown(table)
                markdown_lines.extend(table_md)
                markdown_lines.append("")

            # Write to file
            markdown_content = '\n'.join(markdown_lines)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)

            return output_path

        except Exception as e:
            raise RuntimeError(f"python-docx conversion failed: {e}")

    def _build_markdown_from_docx2python(self, content) -> str:
        """Build markdown content from docx2python extracted content."""
        markdown_lines = []

        # Add document body
        if hasattr(content, 'body') and content.body:
            markdown_lines.extend(self._process_docx2python_body(content.body))

        # Add footnotes if available
        if hasattr(content, 'footnotes') and content.footnotes:
            markdown_lines.append("\n## Footnotes\n")
            markdown_lines.extend(
                self._process_docx2python_footnotes(content.footnotes))

        return '\n'.join(markdown_lines)

    def _process_docx2python_body(self, body) -> list:
        """Process document body from docx2python."""
        lines = []

        # docx2python returns nested lists/tuples representing document structure
        # We need to flatten and format this appropriately
        def process_element(element, level=0):
            if isinstance(element, (list, tuple)):
                for item in element:
                    process_element(item, level + 1)
            elif isinstance(element, str) and element.strip():
                # Clean up the text
                text = element.strip()
                if text:
                    lines.append(text)
                    lines.append("")  # Add spacing

        process_element(body)
        return lines

    def _process_docx2python_footnotes(self, footnotes) -> list:
        """Process footnotes from docx2python."""
        lines = []
        for i, footnote in enumerate(footnotes, 1):
            if isinstance(footnote, str) and footnote.strip():
                lines.append(f"{i}. {footnote.strip()}")
        return lines

    def _convert_table_to_markdown(self, table) -> list:
        """Convert Word table to Markdown format."""
        lines = []
        table_format = self.markdown_options.get('table_format', 'github')

        # Extract table data
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Clean cell text
                cell_text = cell.text.strip().replace('\n', ' ').replace('|', '\\|')
                row_data.append(cell_text)
            table_data.append(row_data)

        if not table_data:
            return lines

        # Generate markdown table
        if table_format == 'github':
            # Header row
            header = table_data[0]
            lines.append('| ' + ' | '.join(header) + ' |')
            lines.append('| ' + ' | '.join(['---'] * len(header)) + ' |')

            # Data rows
            for row in table_data[1:]:
                # Ensure row has same number of columns as header
                while len(row) < len(header):
                    row.append('')
                lines.append('| ' + ' | '.join(row[:len(header)]) + ' |')

        return lines

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        try:
            if 'Heading' in style_name:
                parts = style_name.split()
                if len(parts) > 1 and parts[1].isdigit():
                    return min(int(parts[1]), 6)  # Max heading level is 6
            return 1
        except:
            return 1

    def _post_process_markdown(self, markdown_path: str):
        """Post-process generated markdown for better formatting."""
        try:
            with open(markdown_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # Clean up common formatting issues
            lines = content.split('\n')
            processed_lines = []

            for i, line in enumerate(lines):
                # Remove excessive blank lines
                if line.strip() == '' and i > 0 and processed_lines and processed_lines[-1].strip() == '':
                    continue

                # Clean up table formatting
                if '|' in line and line.strip().startswith('|'):
                    # Ensure proper spacing around pipes
                    line = line.replace('|', ' | ').replace('  |  ', ' | ')
                    line = line.strip()
                    if not line.startswith('|'):
                        line = '| ' + line
                    if not line.endswith('|'):
                        line = line + ' |'

                processed_lines.append(line)

            # Write back the processed content
            with open(markdown_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(processed_lines))

        except Exception as e:
            self.logger.warning(f"Post-processing failed: {e}")

    def is_pandoc_available(self) -> bool:
        """Check if pandoc is available."""
        try:
            pypandoc.get_pandoc_version()
            return True
        except Exception:
            return False

    def get_conversion_info(self) -> Dict[str, Any]:
        """Get information about available conversion methods."""
        return {
            'pandoc_available': self.is_pandoc_available(),
            'docx2python_available': True,  # Always available if installed
            'python_docx_available': True,   # Always available if installed
            'recommended_strategy': 'pandoc' if self.is_pandoc_available() else 'docx2python'
        }

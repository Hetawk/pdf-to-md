"""Word document processing and content extraction."""

import logging
from pathlib import Path
from typing import Dict, Any, List, Optional
import docx
from docx.document import Document
import docx2python


class WordDocumentProcessor:
    """Process Word documents and extract structured content."""

    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.logger = logging.getLogger(__name__)

    def process_document(self, word_path: str) -> Dict[str, Any]:
        """
        Process Word document and extract structured content.

        Args:
            word_path: Path to Word document

        Returns:
            Dictionary containing structured document content
        """
        word_path = Path(word_path)

        if not word_path.exists():
            raise FileNotFoundError(f"Word document not found: {word_path}")

        self.logger.info(f"Processing Word document: {word_path.name}")

        # Extract content using multiple strategies
        content = {
            'metadata': self._extract_metadata(word_path),
            'text_content': self._extract_text_content(word_path),
            'tables': self._extract_tables(word_path),
            'images': self._extract_images(word_path),
            'styles': self._extract_styles(word_path),
            'structure': self._extract_structure(word_path)
        }

        return content

    def _extract_metadata(self, word_path: Path) -> Dict[str, Any]:
        """Extract document metadata."""
        try:
            doc = docx.Document(str(word_path))
            core_props = doc.core_properties

            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'created': core_props.created.isoformat() if core_props.created else None,
                'modified': core_props.modified.isoformat() if core_props.modified else None,
                'revision': core_props.revision or 0,
                'category': core_props.category or '',
                'comments': core_props.comments or ''
            }

            return metadata
        except Exception as e:
            self.logger.warning(f"Failed to extract metadata: {e}")
            return {}

    def _extract_text_content(self, word_path: Path) -> Dict[str, Any]:
        """Extract text content with formatting."""
        try:
            # Use docx2python for structured extraction
            content = docx2python.docx2python(str(word_path))

            return {
                'body': content.body,
                'header': content.header,
                'footer': content.footer,
                'footnotes': content.footnotes,
                'endnotes': content.endnotes,
                'document_xml': content.document_xml if hasattr(content, 'document_xml') else None
            }
        except Exception as e:
            self.logger.warning(f"Failed to extract text content: {e}")
            return {}

    def _extract_tables(self, word_path: Path) -> List[Dict[str, Any]]:
        """Extract tables with formatting preservation."""
        try:
            doc = docx.Document(str(word_path))
            tables = []

            for i, table in enumerate(doc.tables):
                table_data = {
                    'index': i,
                    'rows': [],
                    'style': table.style.name if table.style else None
                }

                for row_idx, row in enumerate(table.rows):
                    row_data = {
                        'index': row_idx,
                        'cells': []
                    }

                    for cell_idx, cell in enumerate(row.cells):
                        cell_data = {
                            'index': cell_idx,
                            'text': cell.text,
                            'paragraphs': [p.text for p in cell.paragraphs]
                        }
                        row_data['cells'].append(cell_data)

                    table_data['rows'].append(row_data)

                tables.append(table_data)

            return tables
        except Exception as e:
            self.logger.warning(f"Failed to extract tables: {e}")
            return []

    def _extract_images(self, word_path: Path) -> List[Dict[str, Any]]:
        """Extract image information."""
        try:
            doc = docx.Document(str(word_path))
            images = []

            # Get all inline shapes (images)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_data = {
                        'relationship_id': rel.rId,
                        'target': rel.target_ref,
                        'content_type': rel.target_part.content_type if hasattr(rel, 'target_part') else None
                    }
                    images.append(image_data)

            return images
        except Exception as e:
            self.logger.warning(f"Failed to extract images: {e}")
            return []

    def _extract_styles(self, word_path: Path) -> Dict[str, Any]:
        """Extract document styles."""
        try:
            doc = docx.Document(str(word_path))
            styles = {
                'paragraph_styles': [],
                'character_styles': [],
                'table_styles': []
            }

            for style in doc.styles:
                style_info = {
                    'name': style.name,
                    'type': str(style.type),
                    'builtin': style.builtin
                }

                if style.type == 1:  # WD_STYLE_TYPE.PARAGRAPH
                    styles['paragraph_styles'].append(style_info)
                elif style.type == 2:  # WD_STYLE_TYPE.CHARACTER
                    styles['character_styles'].append(style_info)
                elif style.type == 3:  # WD_STYLE_TYPE.TABLE
                    styles['table_styles'].append(style_info)

            return styles
        except Exception as e:
            self.logger.warning(f"Failed to extract styles: {e}")
            return {}

    def _extract_structure(self, word_path: Path) -> Dict[str, Any]:
        """Extract document structure (headings, sections)."""
        try:
            doc = docx.Document(str(word_path))
            structure = {
                'headings': [],
                'sections': len(doc.sections),
                'paragraphs': len(doc.paragraphs)
            }

            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.style.name.startswith('Heading'):
                    heading = {
                        'index': i,
                        'text': paragraph.text,
                        'style': paragraph.style.name,
                        'level': self._get_heading_level(paragraph.style.name)
                    }
                    structure['headings'].append(heading)

            return structure
        except Exception as e:
            self.logger.warning(f"Failed to extract structure: {e}")
            return {}

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        try:
            if 'Heading' in style_name:
                # Extract number from style name like "Heading 1"
                parts = style_name.split()
                if len(parts) > 1 and parts[1].isdigit():
                    return int(parts[1])
            return 0
        except:
            return 0
